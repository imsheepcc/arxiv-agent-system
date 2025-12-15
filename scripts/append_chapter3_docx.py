from pathlib import Path

from docx import Document


def main():
    project_root = Path(__file__).resolve().parent.parent
    docx_path = project_root / "report_section0_1_chapter1_en.docx"

    doc = Document(docx_path)

    doc.add_heading("3 System Design (System Design)", level=1)
    doc.add_paragraph(
        """This chapter describes the overall structure and key data flows of the Multi-Agent Code Generation System """
        """from the perspective of system implementation. The core idea is to transform natural-language requirements """
        """into executable tasks via an \"orchestrator + multi-role agents + tooling layer\" approach, then write the """
        """generated results into the actual project directory through tool calls, ultimately producing a runnable """
        """project artifact. To ensure the process is traceable and recoverable, the system persists plans, task states, """
        """generated files, and evaluation conclusions into state files under the output directory during execution, """
        """supporting review and subsequent iterations."""
    )

    doc.add_heading("3.1 Overall Architecture and Layering", level=1)
    doc.add_paragraph(
        """From the project structure perspective, the system can be summarized into five collaborative layers. The entry """
        """layer is responsible for startup and argument parsing and completes the necessary initialization. To ensure """
        """reproducible experiments and enable easy switching across runtime environments, key choices are exposed as """
        """command-line arguments—for example, selecting the model provider and model version via `--provider` and """
        """`--model`, demonstrating the workflow under no-key or offline conditions via `--mock`, and uniformly specifying """
        """the output location for generated artifacts via `--output-dir`. The orchestration layer organizes a run into a """
        """phased workflow of \"planning—generation—evaluation\" and is responsible for dependency checks, task ordering, """
        """and result aggregation during execution, thereby splitting complex requirements into manageable sub-goals. The """
        """agent layer undertakes responsibilities such as requirement decomposition, implementation, and quality review, """
        """and collaborates incrementally under shared context constraints to reduce repeated generation and information """
        """drift. The tooling layer provides agents with controlled model invocation and file read/write capabilities so """
        """that generated content can be stably written into the output directory and validated by reading at any time. The """
        """state layer persists phased progress and key conclusions, supporting interruption recovery and process review, and """
        """ensuring the system does not rely on transient context."""
    )

    doc.add_heading("3.2 Core Execution Flow (Phase-based Workflow)", level=1)
    doc.add_paragraph(
        """The system's core execution is organized in a phase-based manner and can be summarized as three steps: """
        """planning—generation—evaluation. First, in the planning phase, the orchestrator hands the original requirement """
        """to the planning agent for analysis and produces a structured project plan that clarifies target features, """
        """recommended file organization, and an executable task list. This plan is recorded into the state file as the """
        """reference baseline for the subsequent generation phase to avoid deviating from the requirement during execution. """
        """Next, in the code generation phase, the orchestrator advances item by item according to task dependencies and """
        """priorities, and supplies the code generation agent with completed tasks and existing artifacts as context, enabling """
        """it to incrementally improve project files and ultimately form runnable front-end pages, static assets, and data """
        """files in the output directory. Finally, in the evaluation phase, the system checks key generated files against the """
        """requirements, provides a comprehensive evaluation from the perspectives of functional completeness, code quality, """
        """and user experience, and writes the score and improvement suggestions back to the state record (e.g., """
        """`outputs/state/state.json`) to support subsequent iterations."""
    )

    doc.add_heading("3.3 Agent Design (Responsibilities and Interfaces)", level=1)
    doc.add_paragraph(
        """The agent design of this system follows a planning logic of \"clear division of labor, aligned information, and """
        """iterative improvement\": first, a role responsible for global understanding decomposes the requirements into """
        """executable units; then, a role responsible for implementation completes tasks one by one; finally, a role """
        """responsible for quality control checks the overall artifacts for consistency and proposes refinements. Such role """
        """separation decomposes complex problems into a series of more controllable decision points and constrains the input """
        """and output of each step within clear boundaries, thereby reducing common omissions and mismatches when generating an """
        """entire project in one shot."""
    )
    doc.add_paragraph(
        """Specifically, the core value of the planning agent is not \"writing code\" but rewriting natural-language requirements """
        """into a structured plan: clarifying what the target deliverables are, which files and modules are needed, the order """
        """of implementing features, and how dependencies should be arranged. The task list formed in the planning phase is """
        """equivalent to a roadmap for the subsequent execution phase; based on it, the orchestrator performs task ordering and """
        """dependency checks, preventing unreasonable sequences such as rendering pages before generating data files, or finishing """
        """page details before producing shared styles."""
    )
    doc.add_paragraph(
        """The code generation agent undertakes the execution responsibility of \"advancing by tasks\". Its focus is to treat each """
        """task as a small delivery: given a target file and context constraints, it prioritizes using tools to actually write the """
        """content into the project directory and converges toward a runnable implementation through multi-round interaction. """
        """Compared with outputting a large amount of code in one go, this task-oriented execution makes scope control easier, """
        """reduces rework, and allows the orchestrator to feed new artifacts and state back into subsequent tasks after each task """
        """is completed, thereby forming an incremental project construction process."""
    )
    doc.add_paragraph(
        """The evaluation agent is responsible for aligning execution results back with the original requirements. It focuses on """
        """reviewing whether \"the produced artifacts are usable, complete, and meet user experience expectations\" and converts """
        """issues into actionable improvement suggestions as inputs for the next iteration. By separating evaluation from the """
        """generation stage, the system can obtain more objective quality feedback without interrupting the execution rhythm, """
        """and it also makes it easier to introduce stronger validation methods such as automated testing and static analysis in """
        """the future."""
    )

    doc.add_heading("3.4 Communication Protocol (Message Schema & Traceability)", level=1)
    doc.add_paragraph(
        """To make multi-agent collaboration traceable and auditable, the system adopts a unified communication protocol to """
        """describe \"who sent what to whom and when\". During execution, the orchestrator encodes key interactions—such as """
        """planning requests, task dispatch, task completion acknowledgements, evaluation requests, and evaluation reports—into """
        """structured messages and records these messages into each agent's conversation history. The significance is that even """
        """when different agents focus on different concerns at different stages, the system can archive the interaction process """
        """using consistent data structures, thereby supporting review and error localization."""
    )
    doc.add_paragraph(
        """From a process perspective, protocol messages are like a \"ledger\" of collaboration: in the planning phase, requests """
        """and responses bind requirements to plans; in the execution phase, task assignment and result return link each output to """
        """its corresponding task; in the evaluation phase, evaluation requests and reports solidify final conclusions and """
        """improvement suggestions. Because message records align with phase boundaries, when exceptions occur (e.g., incomplete """
        """planning output, a failed task generation, or critical defects found in evaluation), the system can quickly locate which """
        """stage of the collaboration chain the issue happened in and then take strategies such as retrying, rolling back, or """
        """replanning."""
    )

    doc.add_heading("3.5 State and Memory (State & Memory)", level=1)
    doc.add_paragraph(
        """Execution in a multi-agent system often spans multiple steps and multiple model calls. Without persistence, any """
        """interruption or error may lead to context loss, repeated generation, and irreproducible results. Therefore, the system """
        """writes project plans, execution progress, phased artifacts, and evaluation conclusions into state files, and also saves """
        """the conversation history and thought records accumulated by agents during execution, thereby turning \"runtime context\" """
        """into long-term \"process evidence\"."""
    )
    doc.add_paragraph(
        """At the execution level, state persistence directly serves orchestration logic: the orchestrator can obtain completed """
        """tasks and generated files from the state and then provide more accurate context when dispatching new tasks, preventing """
        """agents from recreating the same file or ignoring existing implementations. When execution needs to be resumed, the system """
        """can continue from the last saved state rather than starting over from scratch. Meanwhile, the system also maintains short-term """
        """memory of \"recently generated content\" to strengthen context alignment, enabling subsequent tasks to make incremental adjustments """
        """based on the latest artifacts and improving overall engineering consistency and coherence."""
    )

    doc.add_heading("3.6 Tooling Layer and External Dependencies", level=1)
    doc.add_paragraph(
        """The key reason the system can transform \"model-generated text\" into a \"real runnable project\" is that the tooling layer """
        """provides agents with controlled execution capabilities and abstracts external dependencies and engineering operations into stable """
        """interfaces. For agents, the tooling layer functions as an actionable \"action space\": after the planning phase determines required """
        """files and features, the generation phase does not need to write everything at once; instead, through multi-round interaction, it can """
        """incrementally call tools to create files, write content, and read necessary context, making the execution process more controllable and """
        """easier to pinpoint failure points."""
    )
    doc.add_paragraph(
        """For language model integration, the system uses a unified LLM client to encapsulate differences among model providers so that the upper """
        """orchestration logic can drive different backends using the same message and tool definitions. When switching provider or model, changes """
        """mainly occur in request sending and authentication configuration and do not affect the overall three-phase flow of \"planning—generation—evaluation\". """
        """At the same time, the system provides a fallback strategy for scenarios without an API key: a mock mode keeps the workflow demonstrable and """
        """debuggable. This is particularly important for reproducibility in project environments because it decouples \"external service availability\" from \"system architecture correctness\"."""
    )
    doc.add_paragraph(
        """In terms of engineering operations, file tools restrict actions such as writing files, reading files, and listing directories to within the """
        """output directory. This reduces the risk of misoperations and also gives the generated results a clear boundary: all artifacts are centralized """
        """under the outputs directory, facilitating packaging, demonstration, and testing. Configuration files centralize the management of runtime parameters, """
        """allowing users to control default models, alternative models, and orchestrator parameters via a small number of options, thereby enabling quick switching """
        """between cost and effectiveness across different scenarios."""
    )

    doc.add_heading("3.7 Structure of Generated Artifacts (Generated Artifacts)", level=1)
    doc.add_paragraph(
        """After execution completes, the system generates a well-structured static website project under the output directory. Its organization remains consistent """
        """with the planned file structure produced in the planning phase, reflecting the execution characteristic of \"plan-driven implementation\". Typical artifacts include """
        """page files such as the home page, category pages, and paper detail pages, as well as directories of static assets for unified styling and interaction logic. """
        """In addition, it includes JSON data files that carry paper metadata, usage documentation explaining how to run and view the site locally, and state files. With """
        """this organization, the artifacts can be run and displayed via a simple HTTP server and can also serve as a foundation for subsequent extensions (e.g., adding search, """
        """pagination, and real data ingestion)."""
    )
    doc.add_paragraph(
        """From the runtime data-flow perspective, front-end pages do not hardcode data into HTML; instead, scripts load JSON data in the browser and then render it into """
        """category navigation, paper lists, and detail page content. The advantage is decoupling data from presentation logic: when the data source is updated or extended, """
        """replacing or updating the data file alone can drive changes in page rendering without rewriting the page structure. Meanwhile, interactive features such as citation """
        """generation and copying can also be reused within unified script logic, thereby ensuring consistent experience across different pages."""
    )

    doc.save(docx_path)
    print(f"Appended Chapter 3 to {docx_path}")


if __name__ == "__main__":
    main()
