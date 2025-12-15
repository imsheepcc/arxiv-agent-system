from pathlib import Path

from docx import Document


def main():
    project_root = Path(__file__).resolve().parent.parent
    docx_path = project_root / "report_section0_1_chapter1_en.docx"

    doc = Document(docx_path)

    doc.add_heading("4 Key Challenges and Solutions (Key Challenges and Solutions)", level=1)
    doc.add_paragraph(
        """This chapter summarizes the main challenges encountered by the system in the end-to-end automation process """
        """from natural language to a runnable project, and explains the coping strategies adopted in the current """
        """implementation as well as possible directions for improvement. Because the system is built around multi-agent """
        """collaboration and tool invocation, the challenges are often concentrated in three aspects: whether structured """
        """expression is reliable, whether cross-agent collaboration is consistent, and whether external dependencies and """
        """runtime environments affect reproducibility."""
    )

    doc.add_heading("4.1 Structured Representation from Requirements to Executable Tasks", level=1)
    doc.add_paragraph(
        """The first hurdle for a multi-agent system is converting natural-language requirements into a structured plan """
        """that the orchestrator can execute. In real runs, model outputs may include explanatory text, formatting deviations, """
        """or incomplete JSON fragments, causing the plan to be parsed unstably and thereby affecting subsequent task dispatch. """
        """To mitigate this risk, the planning phase adopts a strategy of \"as structured as possible + degradable on failure\": """
        """it prioritizes extracting the plan JSON from the model output (for example, by locating the outermost brace range """
        """and slicing it). Once parsing fails, a fallback plan is triggered to ensure the system can still produce a minimal """
        """runnable set of files and complete the basic workflow. In this way, the system absorbs uncertainty early in the planning """
        """phase and avoids letting parsing failures propagate across the entire execution chain."""
    )

    doc.add_heading("4.2 Consistency of Multi-Agent Collaboration and Context Alignment", level=1)
    doc.add_paragraph(
        """In a multi-agent division-of-labor setting, consistency issues typically manifest as information loss, repeated task """
        """execution, or generation order confusion due to improper dependency handling. To address this, the system places """
        """\"unified scheduling\" at the orchestrator layer: the orchestrator organizes execution order according to task dependencies """
        """and priorities, and injects necessary context each time it dispatches a task—for example, the list of completed tasks """
        """and generated files (such as `completed_tasks` and `created_files`), as well as summaries of recent artifacts (recent """
        """files/tasks)—so that the execution agent can perform incremental generation based on the latest state. Meanwhile, the system """
        """solidifies phased results through state persistence to form a \"single source of truth\", reducing the probability that different """
        """agents develop inconsistent understandings of the project state."""
    )

    doc.add_heading("4.3 Reliability of Tool Invocation (Function Calling Robustness)", level=1)
    doc.add_paragraph(
        """Tool invocation enhances \"grounding\" capability but also introduces new uncertainty: the model may repeatedly trigger the """
        """same tool call, or produce redundant writes across multiple rounds of interaction, which ultimately appears as duplicate state """
        """records (for example, duplicates in `created_files` or `files_created` in task results). The current implementation's basic """
        """strategy is to advance the dialogue using structured tool return results, enabling the orchestrator and agents to receive verifiable """
        """feedback after each tool execution and proceed accordingly. For more stable engineering deployment, future work can introduce idempotency """
        """and deduplication mechanisms on the orchestrator side—for example, converting file lists to sets before persisting them, or detecting repeated """
        """calls based on target path and write mode before executing tools—thereby turning \"duplicate writes\" from a model-behavior issue into a controllable """
        """system-level constraint."""
    )

    doc.add_heading("4.4 Provider Differences and Portability (Provider Compatibility)", level=1)
    doc.add_paragraph(
        """Different model providers vary in API protocols, message formats, and representations of tool invocation, which directly impacts the portability """
        """and long-term maintenance cost of multi-agent systems. This system adopts a unified client abstraction to shield these differences: the upper """
        """layer consistently drives the model using a unified message structure and a `chat`-style invocation, while the lower layer selects different request/response """
        """adaptation paths according to the provider and performs tool-format conversion when needed (for example, converting a generic tool definition into a schema """
        """accepted by a specific provider). With this encapsulation, \"switching provider/model\" becomes primarily a runtime parameter choice rather than an architectural """
        """rewrite, improving the system's ability to be reused and extended across environments."""
    )

    doc.add_heading("4.5 Reproducibility and Debuggability", level=1)
    doc.add_paragraph(
        """A common pain point of automated generation systems is the difficulty of explaining \"why it generated this way\", and once results deviate from expectations, """
        """it can be hard to locate the cause. To address this, the system emphasizes process tracing: on the one hand, it records key phases, task dispatch, tool """
        """execution, and evaluation conclusions into log files via unified logging (e.g., `logs/agent_system_*.log`); on the other hand, it lets agents record thought """
        """summaries at key points (such as `thoughts`), and, together with protocolized messages, archives interactions including planning requests, task assignments, and """
        """evaluation reports in structured form. As a result, debugging no longer relies solely on comparing final artifacts but can progressively replay and locate issues along """
        """the chain of \"plan—tasks—tool calls—artifacts—evaluation\"."""
    )

    doc.add_heading("4.6 Data Authenticity and External Interface Dependencies (arXiv API)", level=1)
    doc.add_paragraph(
        """In data-driven projects that target real data, reliance on external interfaces is an unavoidable challenge: the requirement emphasizes that the arXiv API must be """
        """called to obtain real papers, but under offline conditions or without an API key, the system cannot guarantee data authenticity and timeliness. To handle this, the """
        """system adopts a compromise strategy of \"optional capability + runnable by degradation\": when external tools are available, it enables the real-data fetching workflow; """
        """when they are unavailable, it allows using sample data so that the website can still run and be demonstrated, while explicitly stating this limitation in the report and """
        """classifying it as a future improvement direction. A further engineering approach is to introduce dedicated data fetching scripts and caching mechanisms, removing online interface """
        """calls from the critical runtime path so that the generation workflow remains controllable and reproducible under unstable network conditions."""
    )

    doc.save(docx_path)
    print(f"Appended Chapter 4 to {docx_path}")


if __name__ == "__main__":
    main()
