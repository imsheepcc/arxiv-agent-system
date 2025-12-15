from pathlib import Path

from docx import Document


def main():
    project_root = Path(__file__).resolve().parent.parent
    docx_path = project_root / "report_section0_1_chapter1_en.docx"

    doc = Document(docx_path)

    doc.add_heading("7 Conclusion (Conclusion)", level=1)
    doc.add_paragraph(
        """In summary, this project verifies the effectiveness of the multi-agent collaboration paradigm for the end-to-end task """
        """of \"from natural-language requirements to runnable project artifacts\". By splitting the development workflow into three """
        """phases—planning, generation, and evaluation—the system can maintain a relatively clear execution path under complex """
        """requirements and transform model outputs into real file artifacts through tool invocation, enabling the final results to be """
        """run and demonstrated directly. Practice shows that this division-of-labor architecture is easier to scope and control than """
        """single-shot generation, and it also makes it easier to locate issues and iterate when problems occur."""
    )
    doc.add_paragraph(
        """Meanwhile, the unified communication protocol and state persistence mechanism provide the system with traceable and reproducible """
        """process records, allowing the collaboration process to be audited and replayed and reducing the uncertainty of \"black-box generation\". """
        """Although there is still room for improvement in areas such as real data acquisition, tool idempotency, and automated verification, """
        """the overall design has laid a stable foundation for subsequent extensions and can support gradual evolution under stricter engineering requirements."""
    )

    doc.add_heading("8 Appendix (Appendix, Optional)", level=1)
    doc.add_paragraph(
        """This appendix summarizes the most critical file locations and commonly used run commands in the project to help readers quickly locate """
        """implementation details and reproduce experimental results."""
    )

    doc.add_heading("8.1 Key File Index", level=1)

    for line in [
        "main.py",
        "orchestrator/multi_agent_orchestrator.py",
        "agents/base_agent.py",
        "agents/code_agent.py",
        "agents/evaluation_agent.py",
        "tools/llm_client.py",
        "tools/file_tools.py",
        "protocol/message_schema.py",
        "state/state_manager.py",
        "outputs/* (generated results)",
    ]:
        doc.add_paragraph(line)

    doc.add_heading("8.2 Run Commands and Parameter Notes", level=1)
    doc.add_paragraph("bash get_real_data.sh")

    try:
        doc.save(docx_path)
        print(f"Appended Chapters 7 and 8 to {docx_path}")
    except PermissionError:
        alt_path = project_root / "report_section0_1_chapter1_en_with_ch7_8.docx"
        doc.save(alt_path)
        print(f"Permission denied when saving to {docx_path}. Saved to {alt_path} instead.")


if __name__ == "__main__":
    main()
