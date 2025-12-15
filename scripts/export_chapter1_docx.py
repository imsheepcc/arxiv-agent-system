"""
Generate a DOCX file containing the English translation of Chapter 1
from report_section0_1.txt. The output file is placed in the project root.
"""

from pathlib import Path

from docx import Document


def build_document() -> Document:
    doc = Document()

    sections = [
        ("1 Project Overview", None),
        (
            None,
            (
                "This project is an assignment for the COMP7103C course titled "
                '"Multi-Agent Code Generation System," and it uses the "arXiv CS Daily" '
                "website generation task as the representative use case. The system runs "
                "as a Python project on Windows, aiming to take natural-language "
                "requirements as input and, through collaboration among multiple agents, "
                "automatically deliver a fully structured software project that can run "
                "out of the box. In line with the course requirements, the system focuses "
                "not only on producing code but also on keeping the generation process "
                "traceable, reproducible, and adaptable to different large-model services."
            ),
        ),
        ("1.1 Project Objectives", None),
        (
            None,
            (
                "From the objective perspective, the system is driven by natural-language "
                "requirements. In the default scenario, the requirement is defined as "
                "`ARXIV_PROJECT_REQUIREMENT` in `prompts/system_prompts.py`, which "
                "describes building an arXiv computer science daily-update website with "
                "discipline-based navigation, paper list browsing, paper detail views, "
                "and citation generation. The system outputs a runnable static front-end "
                "project. The generated results are stored in the `outputs/` directory "
                "and include HTML, CSS, JavaScript, and JSON files that hold paper "
                "metadata, allowing users to obtain deployable web artifacts without "
                "manually constructing the project structure."
            ),
        ),
        ("1.2 Method Overview", None),
        (
            None,
            (
                "Methodologically, the system adopts a multi-agent collaboration paradigm "
                "that splits the plan–implement–evaluate pipeline into three clearly "
                "defined roles: the PlanningAgent interprets the requirement and produces "
                "a structured plan with task lists; the CodeGenerationAgent advances "
                "through the tasks to create the corresponding file contents and commits "
                "them to the filesystem via tool calls; the EvaluationAgent reviews the "
                "resulting codebase for quality, functional completeness, and user "
                "experience, then proposes refinements. These agents are orchestrated by "
                "the MultiAgentOrchestrator, which drives them through staged execution "
                "while maintaining shared context and dependencies to mitigate mismatch "
                "risks under complex requirements. The system also offers a tooling layer: "
                "FileTools handles file creation and I/O, and LLMClient abstracts away the "
                "differences among large-model providers while supporting a mock mode "
                "whenever API keys are unavailable or demonstrations are required, thereby "
                "improving usability and portability."
            ),
        ),
        ("1.3 Results Overview", None),
        (
            None,
            (
                "In terms of outcomes, the system ultimately generates a complete website "
                "structure that includes the home page, category pages, and paper detail "
                "pages. JavaScript loads `data/papers.json` to render data, apply filters, "
                "and provide citation generation along with other key interactions. "
                "According to the persisted evaluation information (see "
                "`outputs/state/state.json`), the current generation achieves an overall "
                "score of 88/100 and is marked as passing; the review also highlights "
                "areas for improvement, such as offering more robust error handling and "
                "friendlier user prompts when data retrieval fails. Overall, the system "
                "demonstrates that multi-agent collaboration can automate the workflow "
                "from requirements to engineering deliverables, and it points to clear "
                "directions for extensions like real arXiv data ingestion, stronger "
                "automated testing, and higher generation stability."
            ),
        ),
    ]

    for title, content in sections:
        if title:
            doc.add_heading(title, level=1)
        if content:
            doc.add_paragraph(content)

    return doc


def main():
    project_root = Path(__file__).resolve().parent.parent
    output_path = project_root / "report_section0_1_chapter1_en.docx"

    document = build_document()
    document.save(output_path)
    print(f"Saved Chapter 1 translation to {output_path}")


if __name__ == "__main__":
    main()
