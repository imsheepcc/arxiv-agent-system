from pathlib import Path

from docx import Document


def main():
    project_root = Path(__file__).resolve().parent.parent
    docx_path = project_root / "report_section0_1_chapter1_en.docx"

    doc = Document(docx_path)

    doc.add_heading("6 Reflections (Reflections)", level=1)
    doc.add_paragraph(
        """This chapter summarizes experiences and takeaways based on the system's actual performance in this use case, """
        """clarifies the current limitations, and proposes actionable directions for improvement. The focus of the reflections """
        """is not to repeat implementation details, but to answer three questions: what real benefits the multi-agent architecture """
        """brought in practice, what issues are still difficult to avoid under the current implementation, and how to advance the """
        """system to a more reliable and more realistic state at relatively low cost."""
    )

    doc.add_heading("6.1 Experiences and Takeaways", level=1)
    doc.add_paragraph(
        """First, the multi-agent division of labor makes the execution of complex requirements more controllable. After separating """
        """\"planning—generation—evaluation\", the planning phase can clarify target files, dependencies, and delivery order; the generation """
        """phase can then progress step by step within a controlled scope; and the evaluation phase can check whether results align with the """
        """requirements from a relatively independent perspective. This phased mechanism reduces omissions and structural disorder that are common """
        """in \"one-shot generation of a whole project\", and it also makes failures at each step easier to locate and roll back. Second, protocolized """
        """messages and state persistence significantly improve reproducibility and debuggability: on the one hand, key interactions are recorded as """
        """structured events; on the other hand, execution progress and evaluation conclusions are written into `outputs/state/state.json` to form a """
        """traceable evidence chain, so debugging no longer depends on memory or manual replay but can be performed by locating and comparing logs and state."""
    )

    doc.add_heading("6.2 Limitations", level=1)
    doc.add_paragraph(
        """The main limitations of the current implementation first come from external data and environment dependencies. The requirement emphasizes """
        """that the arXiv API must be called to obtain real paper data, but under offline conditions or without a key, the system can only use sample """
        """data to validate front-end and workflow correctness, which means the key metric of \"data authenticity, scale, and daily updates\" cannot be """
        """fully proven. Second, the idempotency of tool invocation still has room for improvement: during multi-round tool interactions, the model may """
        """trigger duplicate writes, resulting in duplicated state records or multiple overwrites of artifacts; although this may not break final runnability, """
        """it affects the cleanliness and auditability of the process. Finally, the evaluation phase remains subjective: it currently relies more on model-generated """
        """quality opinions and lacks stronger automated tests, static checks, and consistency validations to support a stricter quality bar."""
    )

    doc.add_heading("6.3 Future Work (Future Work)", level=1)
    doc.add_paragraph(
        """For subsequent iterations, enhancements can be made step by step from the data, front-end, and system perspectives. On the data side, the """
        """priority should be to complete the real data fetching and caching pipeline—for example, implementing `scripts/fetch_papers.py` to solidify the arXiv """
        """API fetching process as a repeatable script, and introducing caching and incremental update strategies so that the \"Daily\" property can still be """
        """achieved under unstable network conditions. On the front-end side, loading and error prompts, search, and pagination can be added on top of the existing """
        """page structure to accommodate browsing needs as real data scale grows, while accessibility (e.g., ARIA and keyboard navigation) can be gradually improved """
        """to enhance general user experience. On the system side, idempotency and validation mechanisms should be strengthened: deduplicate process records such as `created_files`, """
        """apply stricter schema validation to planning and evaluation outputs, and progressively incorporate linting, unit tests, or end-to-end checks into the automation workflow """
        """(for example, running tests via command tools when available and writing results back into evaluation), thereby improving from \"looks usable\" to \"provably reliable\"."""
    )

    try:
        doc.save(docx_path)
        print(f"Appended Chapter 6 to {docx_path}")
    except PermissionError:
        alt_path = project_root / "report_section0_1_chapter1_en_with_ch6.docx"
        doc.save(alt_path)
        print(f"Permission denied when saving to {docx_path}. Saved to {alt_path} instead.")


if __name__ == "__main__":
    main()
