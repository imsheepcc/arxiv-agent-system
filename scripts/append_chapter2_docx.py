from pathlib import Path

from docx import Document


def main():
    project_root = Path(__file__).resolve().parent.parent
    docx_path = project_root / "report_section0_1_chapter1_en.docx"

    doc = Document(docx_path)

    doc.add_heading("2 Introduction", level=1)
    doc.add_paragraph(
        """This section first explains the motivation of the project: as large models increasingly become capable of """
        """\"understanding requirements, generating solutions, and producing code,\" how to organize these capabilities """
        """in an engineering manner to form a stable, traceable, and reproducible automated development workflow has """
        """become a key issue in practical work for software engineering and data mining courses. Compared with having """
        """a single model produce an entire codebase in one shot, a multi-agent division of labor can break complex """
        """requirements into manageable subproblems and improve controllability and extensibility through clear """
        """responsibility boundaries."""
    )

    doc.add_heading("2.1 Requirements and Task Definition", level=1)
    doc.add_paragraph(
        """In terms of requirements and task definition, the system uses \"arXiv CS Daily\" as the default test """
        """requirement, aiming to generate a static website that provides discipline-based category navigation, paper """
        """list browsing, paper detail pages, and citation generation. The project offers a unified entry point via """
        """`main.py`, supports selecting different LLM providers and models using `--provider/--model`, can also run """
        """in demonstration mode without an API key via `--mock`, and uses `--output-dir` and `--max-iterations` to """
        """control the output location and iteration limit. These parameterized designs allow the system to run both """
        """under real API environments and remain usable during in-class demonstrations and debugging."""
    )

    doc.add_heading("2.2 Contributions", level=1)
    doc.add_paragraph(
        """The main contributions of this project are reflected in three aspects: first, it implements a reusable """
        """multi-agent orchestration framework, in which `MultiAgentOrchestrator` is responsible for phased """
        """scheduling and dependency management; second, it defines a unified agent communication protocol """
        """(`protocol/message_schema.py`) to facilitate recording and replaying key interactions; third, it persists """
        """plans, execution progress, and agent memories via `StateManager`, thereby enabling state recovery and """
        """cross-phase context sharing, which provides a foundation for subsequent extensions and experimental """
        """comparisons."""
    )

    doc.save(docx_path)
    print(f"Appended Chapter 2 to {docx_path}")


if __name__ == "__main__":
    main()
